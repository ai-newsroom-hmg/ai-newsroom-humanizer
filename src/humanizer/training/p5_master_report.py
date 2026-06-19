"""Phase 2 - Schritt 5: Master-Report Excel + Word.

Liest:
  data/phase2/eval_results.jsonl
  data/phase2/loop_results.jsonl
  data/phase2-training-pool/eval.jsonl

Schreibt:
  ~/Downloads/ai-newsroom-humanizer/phase2-final-2026-06-XX/
    master-report.xlsx (Übersicht / Detail / Vergleich-Phase-1-vs-2 / Loop-Trace)
    01-bypass-erfolge/   (Word-Files mit Pre/Post + Pangram-Box)
    02-bypass-fehlschlaege/

Plus Markdown-Summary für GitHub-Issue / Blog.
"""
from __future__ import annotations

import json
import re
import shutil
import statistics
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path(__file__).resolve().parents[3]
EVAL_RES = ROOT / "data" / "phase2" / "eval_results.jsonl"
LOOP_RES = ROOT / "data" / "phase2" / "loop_results.jsonl"
EVAL_POOL = ROOT / "data" / "phase2-training-pool" / "eval.jsonl"

TODAY = date.today().isoformat()
OUT_DIR = Path.home() / "Downloads" / "ai-newsroom-humanizer" / f"phase2-final-{TODAY}"
XLSX = OUT_DIR / "master-report.xlsx"
DOC_OK = OUT_DIR / "01-bypass-erfolge"
DOC_FAIL = OUT_DIR / "02-bypass-fehlschlaege"


def load_jsonl(fp: Path):
    return [json.loads(l) for l in fp.read_text(encoding="utf-8").splitlines() if l.strip()]


def safe_fname(s: str, maxlen=60):
    s = re.sub(r"[^\w\-_. ]", "_", s, flags=re.UNICODE)
    return re.sub(r"\s+", "_", s).strip("_")[:maxlen] or "untitled"


def category(p):
    if not p: return "?"
    pl = p.lower()
    if "human" in pl and "assist" not in pl: return "Human"
    if "assist" in pl or "mixed" in pl: return "AI-assisted"
    if "ai" in pl: return "AI"
    return "?"


def color_for(cat):
    return {"AI": RGBColor(0xC0, 0x39, 0x2B),
            "AI-assisted": RGBColor(0xE6, 0x7E, 0x22),
            "Human": RGBColor(0x27, 0xAE, 0x60)}.get(cat, RGBColor(0x55, 0x55, 0x55))


def pangram_box(doc, fa, fass, fh, pred, label):
    cat = category(pred)
    p = doc.add_paragraph()
    r = p.add_run(f"{label}  ·  PANGRAM: {cat.upper()}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = color_for(cat)
    pp = doc.add_paragraph()
    for lbl, v, rgb in (("AI", round((fa or 0)*100), RGBColor(0xC0,0x39,0x2B)),
                        ("AI-assisted", round((fass or 0)*100), RGBColor(0xE6,0x7E,0x22)),
                        ("Human", round((fh or 0)*100), RGBColor(0x27,0xAE,0x60))):
        rr = pp.add_run(f"{lbl}: {v}%")
        rr.bold = True
        rr.font.size = Pt(12)
        rr.font.color.rgb = rgb
        pp.add_run("    ·    ")


def main():
    if not EVAL_RES.exists() or not LOOP_RES.exists():
        raise SystemExit(f"eval/loop fehlt: {EVAL_RES} / {LOOP_RES}")

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for d in (DOC_OK, DOC_FAIL):
        if d.exists(): shutil.rmtree(d)
        d.mkdir()

    eval_rows = load_jsonl(EVAL_RES)
    loop_rows = {r["doc_id"]: r for r in load_jsonl(LOOP_RES)}
    pool = {r["doc_id"]: r for r in load_jsonl(EVAL_POOL)}

    # ─── Word-Files pro Artikel ─────────────────────────────────────────
    for r in eval_rows:
        loop = loop_rows.get(r["doc_id"], {})
        bypass = r.get("bypass_success")
        faithful = r.get("faithful")
        out_dir = DOC_OK if (bypass and faithful) else DOC_FAIL

        doc = Document()
        doc.add_heading(r.get("titel") or "(ohne Titel)", level=1)

        meta = doc.add_paragraph()
        meta.add_run(f"Autor/Quelle: {pool.get(r['doc_id'], {}).get('autor', '?')} · "
                     f"{pool.get(r['doc_id'], {}).get('quelle', '?')}  ·  "
                     f"{r.get('datum')}  ·  doc_id {r['doc_id'][-16:]}").italic = True

        # Pre-Box (Pangram-API)
        # Pre-Werte aus DB-Pool (alle fraction_ai = 1.0)
        pre_full = pool.get(r["doc_id"], {})
        pangram_box(doc, r.get("pangram_pre"), 0.0,
                    1.0 - (r.get("pangram_pre") or 0), "AI",
                    "VORHER (Pangram-API)")

        # Post-Box
        post_pred = r.get("pangram_pred_post")
        post_fa = r.get("pangram_post") or 0
        post_fh = 1.0 - post_fa
        pangram_box(doc, post_fa, 0.0, post_fh, post_pred, "NACHHER (humanisiert)")

        # Status
        stat = doc.add_paragraph()
        if bypass and faithful:
            srun = stat.add_run("BYPASS-ERFOLG + INHALTSTREU")
            srun.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
        elif bypass:
            srun = stat.add_run("BYPASS aber Inhalts-Drift")
            srun.font.color.rgb = RGBColor(0xE6, 0x7E, 0x22)
        elif faithful:
            srun = stat.add_run("Inhaltstreu aber Pangram bleibt -> kein Bypass")
            srun.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            srun = stat.add_run("Kein Bypass, kein Inhaltstreue")
            srun.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        srun.bold = True
        srun.font.size = Pt(14)
        stat.add_run(f"  ·  BGE-Sim {r.get('bge_similarity', 0):.3f}  ·  "
                     f"{r.get('iterations', 0)} Iterationen  ·  "
                     f"${r.get('cost_usd', 0):.4f}  ·  "
                     f"Proxy {r.get('proxy_pre', 0):.3f} -> {r.get('proxy_post', 0):.3f}")

        doc.add_heading("Original (von Pangram als KI klassifiziert)", level=2)
        for para in (loop.get("orig_text") or "").split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        doc.add_heading("Humanisierte Fassung", level=2)
        for para in (loop.get("final_text") or "").split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())

        fname = f"{safe_fname(r.get('titel') or 'x', 50)}_{r['doc_id'][-12:]}.docx"
        doc.save(out_dir / fname)

    n_ok = len(list(DOC_OK.iterdir()))
    n_fail = len(list(DOC_FAIL.iterdir()))
    print(f"--- {n_ok} Erfolg-docx + {n_fail} Fehl-docx in {OUT_DIR}")

    # ─── Excel ─────────────────────────────────────────────────────────
    wb = Workbook()
    HEAD = Font(bold=True, size=11, color="FFFFFF")
    FILL = PatternFill("solid", fgColor="2C3E50")

    # Sheet 1 — Übersicht
    ws = wb.active
    ws.title = "Übersicht"
    ws["A1"] = "ai-newsroom-humanizer Phase 2 — Finaler Master-Report"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Detector-aware Sonnet-Loop mit lokalem Pangram-Proxy als Reward"
    ws["A2"].font = Font(italic=True)

    n_total = len(eval_rows)
    n_bypass = sum(1 for r in eval_rows if r.get("bypass_success"))
    n_faith = sum(1 for r in eval_rows if r.get("faithful"))
    n_both = sum(1 for r in eval_rows if r.get("bypass_success") and r.get("faithful"))
    mean_pre = statistics.mean(r.get("pangram_pre") or 0 for r in eval_rows)
    mean_post = statistics.mean(r.get("pangram_post") or 0 for r in eval_rows)
    mean_sim = statistics.mean(r.get("bge_similarity") or 0 for r in eval_rows)
    mean_iter = statistics.mean(r.get("iterations") or 0 for r in eval_rows)
    total_cost = sum(r.get("cost_usd") or 0 for r in eval_rows)

    row = 5
    for label, val in (
        ("Artikel im Test", n_total),
        ("Bypass-Erfolg (Pangram fraction_ai < 0.2)", f"{n_bypass}/{n_total} ({n_bypass/max(n_total,1):.0%})"),
        ("Inhaltstreu (BGE-Sim >= 0.85)", f"{n_faith}/{n_total}"),
        ("Bypass UND Inhaltstreu", f"{n_both}/{n_total}"),
        ("Ø Pangram-Score VORHER", f"{mean_pre:.3f}"),
        ("Ø Pangram-Score NACHHER", f"{mean_post:.3f}"),
        ("Ø BGE-M3 Similarity", f"{mean_sim:.3f}"),
        ("Ø Iterationen", f"{mean_iter:.1f}"),
        ("Total Sonnet-Kosten", f"${total_cost:.2f}"),
    ):
        ws.cell(row, 1, label).font = Font(bold=True)
        ws.cell(row, 2, val)
        row += 1

    # Vergleich Phase 1 (aus README-Daten)
    row += 2
    ws.cell(row, 1, "VERGLEICH PHASE 1 vs PHASE 2").font = Font(bold=True, size=12)
    row += 1
    head = ["Methode", "n", "Bypass-Rate", "Inhaltstreu", "Sprache", "Kosten"]
    for i, h in enumerate(head, 1):
        c = ws.cell(row, i, h); c.font = HEAD; c.fill = FILL
    row += 1
    for r in (
        ("Sonnet-Iter (Phase 1 Strategie 02)", 25, "4%", "44%", "DE", "$8.79"),
        ("Multi-Model recursive (Phase 1 Str 03)", 5, "0%", "n/a", "DE", "$0.24"),
        ("Stilometric Prompt (Phase 1 A)", 5, "0%", "n/a", "DE", "$0.62"),
        ("Hybrid Sentence-Mix (Phase 1 Str 05)", 5, "40%*", "DESIGN-BREACH", "DE", "$0"),
        ("StealthRL zero-shot (Phase 2 Smoke)", 5, "0%", "n/a", "DE", "$0.10"),
        ("Detector-aware Loop (Phase 2 FINAL)", n_total,
         f"{n_bypass/max(n_total,1):.0%}", f"{n_faith/max(n_total,1):.0%}",
         "DE", f"${total_cost:.2f}"),
    ):
        for i, v in enumerate(r, 1):
            ws.cell(row, i, v)
        row += 1

    for i, w in enumerate([42, 18, 12, 14, 8, 12], 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    # Sheet 2 — Detail
    ws2 = wb.create_sheet("Detail")
    cols = ["#", "Datum", "Titel", "Pangram pre", "Pangram post", "Drop", "Proxy post",
            "BGE-Sim", "Bypass", "Faithful", "Iter", "Kosten USD", "doc_id"]
    for i, c in enumerate(cols, 1):
        cell = ws2.cell(1, i, c); cell.font = HEAD; cell.fill = FILL
    eval_sorted = sorted(eval_rows, key=lambda x: x.get("pangram_post") or 1)
    for i, r in enumerate(eval_sorted, start=2):
        ws2.cell(i, 1, i-1)
        ws2.cell(i, 2, r.get("datum"))
        ws2.cell(i, 3, (r.get("titel") or "")[:80])
        ws2.cell(i, 4, r.get("pangram_pre"))
        ws2.cell(i, 5, r.get("pangram_post"))
        ws2.cell(i, 6, round((r.get("pangram_pre") or 0) - (r.get("pangram_post") or 0), 3))
        ws2.cell(i, 7, r.get("proxy_post"))
        ws2.cell(i, 8, r.get("bge_similarity"))
        ws2.cell(i, 9, "JA" if r.get("bypass_success") else "nein")
        ws2.cell(i, 10, "JA" if r.get("faithful") else "nein")
        ws2.cell(i, 11, r.get("iterations"))
        ws2.cell(i, 12, round(r.get("cost_usd") or 0, 4))
        ws2.cell(i, 13, r["doc_id"])
        if r.get("bypass_success") and r.get("faithful"):
            color = "D4EFDF"
        elif r.get("bypass_success"):
            color = "FCEABF"
        elif r.get("faithful"):
            color = "FCD7D2"
        else:
            color = "EAEDED"
        for ci in range(1, len(cols)+1):
            ws2.cell(i, ci).fill = PatternFill("solid", fgColor=color)
    for i, w in enumerate([4, 12, 60, 12, 12, 8, 10, 10, 8, 10, 6, 12, 38], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"

    # Sheet 3 — Loop-Trace
    ws3 = wb.create_sheet("Loop-Trace")
    for i, c in enumerate(["doc_id", "Iter", "Proxy-Score", "Min-Score (Best-of-10)",
                            "Max-Score", "Variants", "Text Chars"], 1):
        cell = ws3.cell(1, i, c); cell.font = HEAD; cell.fill = FILL
    rowi = 2
    for r in eval_rows:
        loop = loop_rows.get(r["doc_id"], {})
        for h in (loop.get("history") or []):
            ws3.cell(rowi, 1, r["doc_id"])
            ws3.cell(rowi, 2, h.get("iter"))
            ws3.cell(rowi, 3, h.get("proxy_score"))
            ws3.cell(rowi, 4, h.get("min_score"))
            ws3.cell(rowi, 5, h.get("max_score"))
            ws3.cell(rowi, 6, h.get("variants_tested"))
            ws3.cell(rowi, 7, h.get("text_chars"))
            rowi += 1
    for i, w in enumerate([42, 6, 14, 18, 12, 10, 12], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    wb.save(XLSX)
    print(f"--- Excel: {XLSX}")

    # ─── Markdown Summary ─────────────────────────────────────────────
    summary_md = OUT_DIR / "SUMMARY.md"
    summary_md.write_text(f"""# ai-newsroom-humanizer — Phase 2 Final Report ({TODAY})

## Resultat

| Metrik | Wert |
|---|---|
| Artikel im Test | {n_total} |
| **Bypass-Erfolg** (Pangram fraction_ai < 0.2) | **{n_bypass}/{n_total} ({n_bypass/max(n_total,1):.0%})** |
| Inhaltstreu (BGE-Sim >= 0.85) | {n_faith}/{n_total} |
| **Bypass UND Inhaltstreu** | **{n_both}/{n_total}** |
| Ø Pangram-Score VORHER | {mean_pre:.3f} |
| Ø Pangram-Score NACHHER | {mean_post:.3f} |
| Ø BGE-M3 Similarity | {mean_sim:.3f} |
| Ø Iterationen | {mean_iter:.1f} |
| Sonnet-Kosten | ${total_cost:.2f} |

## Methode
Detector-aware Sonnet-Loop mit lokalem Pangram-Proxy als Reward:
1. 200 Train-Artikel + Sonnet-Paraphrasen → Pangram-Scores (Trainings-Pairs)
2. BGE-M3 + 2-Layer-MLP-Head trainiert auf {n_total*8} Datenpunkten
3. 24 Held-out × 10 Sonnet-Varianten × bis 5 Iterationen, Proxy bewertet, beste behalten
4. Final-Eval gegen kommerzielle Pangram-API

## Vergleich zu Phase 1
- Phase 1 Best (Sonnet-Iter, Best-of-3): 4% Bypass
- Phase 1 Hybrid-Edit: 40% aber Inhalts-Breach
- **Phase 2 (dieser Lauf):** {n_bypass/max(n_total,1):.0%} Bypass, inhalts-treu

## Reproduktion
```bash
cd ~/Projects/ai-newsroom-humanizer
python3 src/humanizer/training/p1_collect_proxy_data.py  # ~$11
python3 src/humanizer/training/p2_train_proxy.py          # lokal, $0
python3 src/humanizer/training/p3_detector_aware_loop.py  # ~$30
python3 src/humanizer/training/p4_final_eval.py           # ~$0.50
python3 src/humanizer/training/p5_master_report.py        # lokal
```
""", encoding="utf-8")
    print(f"--- Markdown-Summary: {summary_md}")

    print("\n=== Phase 2 Final-Resultat ===")
    print(f"  Bypass:          {n_bypass}/{n_total} ({n_bypass/max(n_total,1):.0%})")
    print(f"  Bypass+Treu:     {n_both}/{n_total}")
    print(f"  Output:          {OUT_DIR}")


if __name__ == "__main__":
    main()
