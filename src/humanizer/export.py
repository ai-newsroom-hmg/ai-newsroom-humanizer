"""Humanize-Phase 03 — Excel + Word-Export Pre/Post.

Excel-Sheets:
  Uebersicht          — Aggregat (% Erfolg, Ø-Iterationen, Ø-fraction_ai-Drop)
  Pre-Post            — alle Artikel mit fraction_ai vor/nach + Faithfulness-Spalte
  Iterations-Trace    — Verlauf pro Artikel + Iteration

Word-Files:
  ~/Downloads/ki-check-humanize-test/01-loop/<doc_id>.docx
    Original + Final mit Pangram-Box pre/post + Faithfulness-Urteil
"""
from __future__ import annotations

import json
import re
import shutil
import statistics
from collections import Counter
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

ROOT = Path.home() / "Projects" / "ki-check"
LOOP = ROOT / "data" / "humanize" / "01_casdorff_loop_results.jsonl"
FAITH = ROOT / "data" / "humanize" / "02_faithfulness.jsonl"

OUT_DIR = Path.home() / "Downloads" / "ki-check-humanize-test"
DOC_DIR = OUT_DIR / "01-loop"
XLSX = OUT_DIR / "humanize-ergebnisse.xlsx"


def load_jsonl(fp):
    return [json.loads(l) for l in fp.read_text(encoding="utf-8").splitlines() if l.strip()]


def safe_fname(s, maxlen=60):
    s = re.sub(r"[^\w\-_. ]", "_", s, flags=re.UNICODE)
    return re.sub(r"\s+", "_", s).strip("_")[:maxlen] or "untitled"


def category(p):
    if not p:
        return "?"
    pl = p.lower()
    if "human" in pl and "assist" not in pl: return "Human"
    if "assist" in pl or "mixed" in pl: return "AI-assisted"
    if "ai" in pl: return "AI"
    return "?"


def color_for(cat):
    return {"AI": RGBColor(0xC0, 0x39, 0x2B),
            "AI-assisted": RGBColor(0xE6, 0x7E, 0x22),
            "Human": RGBColor(0x27, 0xAE, 0x60)}.get(cat, RGBColor(0x55, 0x55, 0x55))


def pangram_box(doc, fraction_ai, fraction_assisted, fraction_human, prediction, label):
    """Schreibt prominente Pangram-Box in eine docx."""
    cat = category(prediction)
    p = doc.add_paragraph()
    r0 = p.add_run(f"{label}  ·  PANGRAM: {cat.upper()}")
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = color_for(cat)

    pp = doc.add_paragraph()
    for lbl, val, rgb in (
        ("AI", round((fraction_ai or 0) * 100),
         RGBColor(0xC0, 0x39, 0x2B)),
        ("AI-assisted", round((fraction_assisted or 0) * 100),
         RGBColor(0xE6, 0x7E, 0x22)),
        ("Human", round((fraction_human or 0) * 100),
         RGBColor(0x27, 0xAE, 0x60)),
    ):
        rr = pp.add_run(f"{lbl}: {val}%")
        rr.bold = True
        rr.font.size = Pt(12)
        rr.font.color.rgb = rgb
        pp.add_run("    ·    ")


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    if DOC_DIR.exists():
        shutil.rmtree(DOC_DIR)
    DOC_DIR.mkdir(parents=True)

    loop_rows = load_jsonl(LOOP)
    faith_rows = {r["doc_id"]: r for r in (load_jsonl(FAITH) if FAITH.exists() else [])}

    # ============== Word-Files ==============
    n_docs = 0
    for r in loop_rows:
        if "error" in r or not r.get("final_text"):
            continue
        f = faith_rows.get(r["doc_id"], {})
        struct = f.get("structural", {})
        judge = f.get("llm_judge", {})

        doc = Document()
        doc.add_heading(r.get("titel", "(ohne Titel)"), level=1)

        meta = doc.add_paragraph()
        meta.add_run(f"Autor: {r.get('autor','?')}").bold = True
        meta.add_run(f"  ·  Tagesspiegel  ·  {r.get('datum','?')}  ·  {r.get('woerter','?')} Wörter Original")
        meta.add_run(f"  ·  doc_id: {r['doc_id'][:20]}…")

        # Pangram-Boxen
        h_pre = r["history"][0]
        h_post = r["history"][-1]
        pangram_box(doc, h_pre["fraction_ai"], h_pre["fraction_ai_assisted"],
                    h_pre["fraction_human"], h_pre["prediction"], "VORHER (Original)")
        pangram_box(doc, h_post["fraction_ai"], h_post["fraction_ai_assisted"],
                    h_post["fraction_human"], h_post["prediction"], "NACHHER (humanisiert)")

        # Erfolg / Iterationen
        ip = doc.add_paragraph()
        success = r.get("success", False)
        rr = ip.add_run(f"Bypass-Resultat: {'ERFOLG' if success else 'TEILERFOLG/FEHLSCHLAG'}")
        rr.bold = True
        rr.font.size = Pt(13)
        rr.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) if success else RGBColor(0xC0, 0x39, 0x2B)
        ip.add_run(f"  ·  Iterationen: {r.get('iterations_run','?')}/{5}  "
                   f"·  Kosten: {r.get('total_cost_usd', 0):.4f} USD")

        # Faithfulness-Block
        if f:
            doc.add_heading("Inhalts-Treue-Pruefung", level=2)
            fp = doc.add_paragraph()
            verdict = judge.get("verdict") or "?"
            fp_run = fp.add_run(f"LLM-Urteil: {verdict}")
            fp_run.bold = True
            fp_run.font.size = Pt(12)
            fp_run.font.color.rgb = {
                "FAITHFUL": RGBColor(0x27, 0xAE, 0x60),
                "MINOR_DRIFT": RGBColor(0xE6, 0x7E, 0x22),
                "SIGNIFICANT_DRIFT": RGBColor(0xC0, 0x39, 0x2B),
                "CONTENT_CHANGED": RGBColor(0xC0, 0x39, 0x2B),
            }.get(verdict, RGBColor(0x55, 0x55, 0x55))
            fp.add_run(f"  ·  Faithfulness-Score: {judge.get('faithfulness_score','?')}")
            if judge.get("comment"):
                doc.add_paragraph(f"Kommentar: {judge['comment']}")

            sp = doc.add_paragraph()
            sp.add_run(
                f"Strukturell: 4-gram-Overlap {struct.get('four_gram_overlap', 0)*100:.1f} % · "
                f"Namen erhalten {struct.get('names_preserved_ratio', 0)*100:.0f} % · "
                f"Zahlen erhalten {struct.get('numbers_preserved_ratio', 0)*100:.0f} %"
            ).italic = True

            if struct.get("names_lost"):
                doc.add_paragraph("Namen verloren: " + ", ".join(struct["names_lost"][:10]))
            if struct.get("numbers_lost"):
                doc.add_paragraph("Zahlen verloren: " + ", ".join(struct["numbers_lost"][:10]))
            if judge.get("missing"):
                pp = doc.add_paragraph()
                pp.add_run("FEHLENDE Aussagen: ").bold = True
                pp.add_run("; ".join(judge["missing"][:5]))
            if judge.get("added"):
                pp = doc.add_paragraph()
                pp.add_run("HINZUGEFÜGTE Aussagen (unerlaubt): ").bold = True
                pp.add_run("; ".join(judge["added"][:5]))

        # Original
        doc.add_heading("Original (Tagesspiegel)", level=2)
        for para in (r.get("original_text") or "").split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        # Humanisierte Version
        doc.add_heading("Humanisierte Version", level=2)
        for para in (r.get("final_text") or "").split("\n"):
            para = para.strip()
            if para:
                doc.add_paragraph(para)

        # Iterations-Trace
        if len(r.get("history", [])) > 1:
            doc.add_heading("Iterations-Verlauf", level=2)
            for h in r["history"]:
                doc.add_paragraph(
                    f"Iter {h['iter']}: fraction_ai={h['fraction_ai']:.3f} "
                    f"({h.get('prediction','?')})"
                )

        fname = f"{safe_fname(r.get('titel','x'), 50)}_{r['doc_id'][-12:]}.docx"
        doc.save(DOC_DIR / fname)
        n_docs += 1

    print(f"--- {n_docs} Pre/Post-.docx in {DOC_DIR}")

    # ============== Excel ==============
    wb = Workbook()
    HEAD = Font(bold=True, size=11, color="FFFFFF")
    FILL = PatternFill("solid", fgColor="2C3E50")

    # ── Sheet 1: Uebersicht ──
    ws = wb.active
    ws.title = "Übersicht"
    ws["A1"] = "Humanize-Test gegen Pangram — Casdorff-Korpus (Tagesspiegel)"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = "Strategie 02: iterativer Sonnet+Pangram-Loop (AuthorMist-light) — Threshold P(AI) < 0.2"
    ws["A2"].font = Font(italic=True, color="555555")

    rows_ok = [r for r in loop_rows if "error" not in r]
    n_total = len(rows_ok)
    n_success = sum(1 for r in rows_ok if r.get("success"))
    avg_iters = statistics.mean(r.get("iterations_run", 0) for r in rows_ok) if n_total else 0
    avg_drop = statistics.mean((r.get("fraction_ai_pre", 1) - r.get("fraction_ai_post", 1))
                                for r in rows_ok) if n_total else 0
    total_cost = sum(r.get("total_cost_usd", 0) for r in rows_ok)

    # Faithfulness Aggregate
    faith_verdicts = Counter()
    avg_4gram = avg_names = avg_nums = 0.0
    n_faith = 0
    for r in rows_ok:
        f = faith_rows.get(r["doc_id"])
        if not f:
            continue
        v = (f.get("llm_judge") or {}).get("verdict")
        if v:
            faith_verdicts[v] += 1
        s = f.get("structural") or {}
        avg_4gram += s.get("four_gram_overlap", 0)
        avg_names += s.get("names_preserved_ratio", 0)
        avg_nums += s.get("numbers_preserved_ratio", 0)
        n_faith += 1
    if n_faith:
        avg_4gram /= n_faith
        avg_names /= n_faith
        avg_nums /= n_faith

    ws["A5"] = "Bypass-Erfolg (Pangram knacken)"
    ws["A5"].font = Font(bold=True, size=12)
    row = 6
    for label, val in (
        ("Artikel im Test", n_total),
        ("Erfolgreich auf P(AI) < 0.2 gebracht", f"{n_success}/{n_total} ({n_success/max(n_total,1):.0%})"),
        ("Ø Iterationen bis Erfolg/Stop", f"{avg_iters:.1f}"),
        ("Ø fraction_ai-Drop", f"{avg_drop:.3f}"),
        ("Total Kosten Humanize", f"{total_cost:.4f} USD"),
    ):
        ws.cell(row, 1, label).font = Font(bold=True)
        ws.cell(row, 2, val)
        row += 1

    row += 1
    ws.cell(row, 1, "Inhalts-Treue (Faithfulness)").font = Font(bold=True, size=12)
    row += 1
    for v in ("FAITHFUL", "MINOR_DRIFT", "SIGNIFICANT_DRIFT", "CONTENT_CHANGED"):
        ws.cell(row, 1, v).font = Font(bold=True)
        ws.cell(row, 2, f"{faith_verdicts.get(v, 0)}/{n_faith}")
        row += 1
    ws.cell(row, 1, "Ø 4-gram-Overlap mit Original").font = Font(bold=True)
    ws.cell(row, 2, f"{avg_4gram*100:.1f} %")
    row += 1
    ws.cell(row, 1, "Ø Namen erhalten").font = Font(bold=True)
    ws.cell(row, 2, f"{avg_names*100:.1f} %")
    row += 1
    ws.cell(row, 1, "Ø Zahlen erhalten").font = Font(bold=True)
    ws.cell(row, 2, f"{avg_nums*100:.1f} %")

    ws.column_dimensions["A"].width = 38
    ws.column_dimensions["B"].width = 28

    # ── Sheet 2: Pre-Post-Detail ──
    ws2 = wb.create_sheet("Pre-Post")
    cols = ["#", "Datum", "Titel", "Wörter Original", "f_ai PRE", "f_ai POST",
            "Drop", "Iterationen", "Erfolg",
            "Faithfulness", "4-gram %", "Namen %", "Zahlen %",
            "Kosten USD", "doc_id"]
    for i, c in enumerate(cols, 1):
        cell = ws2.cell(1, i, c)
        cell.font = HEAD
        cell.fill = FILL
        cell.alignment = Alignment(horizontal="center")

    rowi = 2
    rows_ok.sort(key=lambda r: r.get("fraction_ai_post", 1))
    for i, r in enumerate(rows_ok, 1):
        f = faith_rows.get(r["doc_id"], {})
        s = f.get("structural", {})
        j = f.get("llm_judge", {})
        pre = r.get("fraction_ai_pre", 1)
        post = r.get("fraction_ai_post", 1)
        ws2.cell(rowi, 1, i)
        ws2.cell(rowi, 2, r.get("datum"))
        ws2.cell(rowi, 3, r.get("titel"))
        ws2.cell(rowi, 4, r.get("woerter"))
        ws2.cell(rowi, 5, round(pre, 3))
        ws2.cell(rowi, 6, round(post, 3))
        ws2.cell(rowi, 7, round(pre - post, 3))
        ws2.cell(rowi, 8, r.get("iterations_run"))
        ws2.cell(rowi, 9, "JA" if r.get("success") else "nein")
        ws2.cell(rowi, 10, j.get("verdict", "?"))
        ws2.cell(rowi, 11, round(s.get("four_gram_overlap", 0) * 100, 1))
        ws2.cell(rowi, 12, round(s.get("names_preserved_ratio", 0) * 100, 0))
        ws2.cell(rowi, 13, round(s.get("numbers_preserved_ratio", 0) * 100, 0))
        ws2.cell(rowi, 14, round(r.get("total_cost_usd", 0), 4))
        ws2.cell(rowi, 15, r["doc_id"])

        # Farb-Codierung: Erfolg & Faithfulness
        success = r.get("success", False)
        v = j.get("verdict")
        if success and v == "FAITHFUL":
            color = "D4EFDF"  # gruen
        elif success and v in ("MINOR_DRIFT", None):
            color = "FCEABF"  # gelb
        elif success and v in ("SIGNIFICANT_DRIFT", "CONTENT_CHANGED"):
            color = "FCD7D2"  # rot
        elif not success:
            color = "EAEDED"  # grau
        else:
            color = None
        if color:
            for ci in range(1, len(cols) + 1):
                ws2.cell(rowi, ci).fill = PatternFill("solid", fgColor=color)
        rowi += 1

    for i, w in enumerate([4, 12, 50, 10, 10, 10, 8, 8, 8, 18, 10, 10, 10, 12, 38], 1):
        ws2.column_dimensions[get_column_letter(i)].width = w
    ws2.freeze_panes = "A2"
    ws2.auto_filter.ref = f"A1:{get_column_letter(len(cols))}{rowi-1}"

    # ── Sheet 3: Iterations-Trace ──
    ws3 = wb.create_sheet("Iterations-Trace")
    for i, c in enumerate(["doc_id", "Iter", "f_ai", "f_ai_assisted", "f_human",
                            "prediction", "text_chars", "Quelle"], 1):
        cell = ws3.cell(1, i, c)
        cell.font = HEAD
        cell.fill = FILL
    rowi = 2
    for r in rows_ok:
        for h in r.get("history", []):
            ws3.cell(rowi, 1, r["doc_id"])
            ws3.cell(rowi, 2, h["iter"])
            ws3.cell(rowi, 3, round(h["fraction_ai"], 3))
            ws3.cell(rowi, 4, round(h["fraction_ai_assisted"], 3))
            ws3.cell(rowi, 5, round(h["fraction_human"], 3))
            ws3.cell(rowi, 6, h.get("prediction"))
            ws3.cell(rowi, 7, h.get("text_chars"))
            ws3.cell(rowi, 8, h.get("source"))
            rowi += 1
    for i, w in enumerate([42, 6, 8, 12, 9, 14, 10, 14], 1):
        ws3.column_dimensions[get_column_letter(i)].width = w
    ws3.freeze_panes = "A2"

    wb.save(XLSX)
    print(f"--- Excel: {XLSX}")
    print(f"\n=== Bypass-Resultat ===")
    print(f"  Erfolg: {n_success}/{n_total} ({n_success/max(n_total,1):.0%})")
    print(f"  Ø Iterationen: {avg_iters:.1f}")
    print(f"  Ø fraction_ai-Drop: {avg_drop:.3f}")
    print(f"  Faithfulness: {dict(faith_verdicts)}")


if __name__ == "__main__":
    main()
